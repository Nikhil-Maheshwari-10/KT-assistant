"""
app/api/routers/documents.py
KT document generation and export endpoints — fully self-contained.

  POST /api/sessions/{session_id}/document/generate → generate markdown
  GET  /api/sessions/{session_id}/document          → get last generated doc
  GET  /api/sessions/{session_id}/document/pdf      → stream PDF bytes
  GET  /api/sessions/{session_id}/document/docx     → stream DOCX bytes
"""
from __future__ import annotations

import base64
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import urllib.request
import uuid
from datetime import datetime

from fastapi import APIRouter, Depends
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field

from app.core.logger import logger
from app.core.messages import DOCUMENT_NOT_FOUND, DOCUMENT_NO_CONTENT, DOCUMENT_PDF_ERROR, DOCUMENT_DOCX_ERROR
from app.core.exceptions import NotFoundException, UnprocessableException, AppException
from app.api.deps import get_session_or_404
from app.services.db_service import db_service
from app.services.ai_engine import ai_engine
from app.models.schemas import Session

router = APIRouter(prefix="/api/sessions/{session_id}/document", tags=["Documents"])

# In-memory cache keyed by session_id
_document_cache: dict[str, str] = {}
MIN_DIAGRAM_BYTES = 3000


# ---------------------------------------------------------------------------
# Response model (local)
# ---------------------------------------------------------------------------

class DocumentResponse(BaseModel):
    session_id: str
    markdown: str
    generated_at: datetime = Field(default_factory=datetime.now)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _download_mermaid_image(code: str):
    """Download a Mermaid diagram PNG from mermaid.ink. Returns (bytes, mime) or None."""
    state = {"code": code, "mermaid": {"theme": "default"}}
    b64 = base64.urlsafe_b64encode(json.dumps(state).encode("utf-8")).decode("utf-8")
    try:
        req = urllib.request.Request(
            f"https://mermaid.ink/img/{b64}",
            headers={"User-Agent": "Mozilla/5.0"},
        )
        with urllib.request.urlopen(req, timeout=15) as resp:
            data = resp.read()
        if data[:4] == b"\x89PNG":
            mime = "image/png"
        elif data[:3] == b"\xff\xd8\xff":
            mime = "image/jpeg"
        else:
            return None
        return (data, mime) if len(data) >= MIN_DIAGRAM_BYTES else None
    except Exception as e:
        logger.error(f"Mermaid download error: {e}")
        return None


# ---------------------------------------------------------------------------
# POST /generate
# ---------------------------------------------------------------------------

@router.post("/generate", response_model=DocumentResponse)
async def generate_document(
    session_id: str,
    session: Session = Depends(get_session_or_404),
):
    """Generates a professional KT markdown document from session knowledge."""
    if not any(t.confidence_score > 0 for t in session.topics):
        raise UnprocessableException(DOCUMENT_NO_CONTENT)

    markdown = ai_engine.generate_final_summary(session)
    markdown = re.sub(r"<[^>]+>", "", markdown)  # strip stray HTML
    _document_cache[session_id] = markdown
    logger.info(f"[DOC] Document generated for session {session_id} ({len(markdown)} chars)")
    return DocumentResponse(session_id=session_id, markdown=markdown)


# ---------------------------------------------------------------------------
# GET /
# ---------------------------------------------------------------------------

@router.get("", response_model=DocumentResponse)
async def get_document(
    session_id: str,
    session: Session = Depends(get_session_or_404),
):
    """Returns the last generated markdown document."""
    markdown = _document_cache.get(session_id)
    if not markdown:
        raise NotFoundException(DOCUMENT_NOT_FOUND)
    return DocumentResponse(session_id=session_id, markdown=markdown)


# ---------------------------------------------------------------------------
# GET /pdf
# ---------------------------------------------------------------------------

@router.get("/pdf")
async def download_pdf(
    session_id: str,
    session: Session = Depends(get_session_or_404),
):
    """Converts the markdown to PDF via Playwright and streams bytes."""
    markdown = _document_cache.get(session_id)
    if not markdown:
        raise NotFoundException(DOCUMENT_NOT_FOUND)

    import markdown as md_lib

    pdf_md = re.sub(
        r"```mermaid\n(.*?)\n```",
        lambda m: f'<div class="mermaid" style="text-align:center;">\n{m.group(1).strip()}\n</div>',
        markdown, flags=re.DOTALL,
    )
    html_body = md_lib.markdown(pdf_md, extensions=["tables", "fenced_code"])
    full_html = f"""<!DOCTYPE html><html><head>
<script type="module">
    import mermaid from 'https://cdn.jsdelivr.net/npm/mermaid@11/dist/mermaid.esm.min.mjs';
    mermaid.initialize({{ startOnLoad: true, theme: 'default' }});
</script>
<style>
body {{ font-family: -apple-system,BlinkMacSystemFont,"Segoe UI",Helvetica,Arial,sans-serif;
       line-height:1.6; color:#24292f; margin:40px auto; max-width:900px; padding:20px; }}
h1,h2,h3,h4 {{ border-bottom:1px solid #d0d7de; padding-bottom:.3em; margin-top:24px; font-weight:600; }}
table {{ border-collapse:collapse; width:100%; margin:16px 0; font-size:14px; }}
th,td {{ border:1px solid #d0d7de; padding:8px 13px; text-align:left; }}
th {{ font-weight:600; background:#f6f8fa; }}
tr:nth-child(2n) {{ background:#f6f8fa; }}
code {{ background:#afb8c133; padding:.2em .4em; border-radius:6px; font-family:monospace; font-size:85%; }}
pre {{ background:#f6f8fa; padding:16px; border-radius:6px; overflow:auto; }}
pre code {{ background:transparent; padding:0; }}
img {{ max-width:100%; display:block; margin:20px auto; }}
</style></head><body>{html_body}</body></html>"""

    tmp_dir = tempfile.mkdtemp()
    html_path = os.path.join(tmp_dir, "doc.html")
    pdf_path = os.path.join(tmp_dir, "doc.pdf")
    try:
        with open(html_path, "w", encoding="utf-8") as f:
            f.write(full_html)
        script_path = os.path.join(
            os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))),
            "scripts", "generate_pdf.py",
        )
        subprocess.run([sys.executable, script_path, html_path, pdf_path], check=True, timeout=120)
        with open(pdf_path, "rb") as f:
            pdf_bytes = f.read()
        return StreamingResponse(iter([pdf_bytes]), media_type="application/pdf",
                                 headers={"Content-Disposition": f'attachment; filename="KT_{session_id[:8]}.pdf"'})
    except subprocess.CalledProcessError:
        raise AppException(message=DOCUMENT_PDF_ERROR, status_code=500)
    except Exception as e:
        raise AppException(message=str(e), status_code=500)
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)


# ---------------------------------------------------------------------------
# GET /docx
# ---------------------------------------------------------------------------

@router.get("/docx")
async def download_docx(
    session_id: str,
    session: Session = Depends(get_session_or_404),
):
    """Converts the markdown to DOCX via Pandoc and streams bytes."""
    markdown = _document_cache.get(session_id)
    if not markdown:
        raise NotFoundException(DOCUMENT_NOT_FOUND)

    tmp_dir = tempfile.mkdtemp()
    try:
        def _mermaid_to_local(match):
            result = _download_mermaid_image(match.group(1).strip())
            if result:
                img_bytes, mime = result
                ext = "png" if mime == "image/png" else "jpg"
                local_path = os.path.join(tmp_dir, f"mermaid_{uuid.uuid4().hex[:8]}.{ext}")
                with open(local_path, "wb") as out:
                    out.write(img_bytes)
                return f"![Diagram]({local_path})"
            return "_[Diagram — could not be rendered]_"

        docx_md = re.sub(r"```mermaid\n(.*?)\n```", _mermaid_to_local, markdown, flags=re.DOTALL)
        md_path = os.path.join(tmp_dir, "doc.md")
        docx_path = os.path.join(tmp_dir, "doc.docx")
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(docx_md)

        proc = subprocess.run(["pandoc", md_path, "-o", docx_path],
                               capture_output=True, text=True, timeout=120)
        if proc.returncode != 0:
            raise AppException(message=DOCUMENT_DOCX_ERROR.format(proc.stderr), status_code=500)

        # Polish tables with python-docx
        try:
            from docx import Document as DocxDocument
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            from docx.shared import Pt

            def _border(cell):
                tc = cell._tc; tcPr = tc.get_or_add_tcPr()
                b = OxmlElement("w:tcBorders")
                for edge in ("top","left","bottom","right","insideH","insideV"):
                    t = OxmlElement(f"w:{edge}")
                    t.set(qn("w:val"),"single"); t.set(qn("w:sz"),"6")
                    t.set(qn("w:space"),"0"); t.set(qn("w:color"),"4A4A4A")
                    b.append(t)
                tcPr.append(b)

            def _pad(cell, pad=120):
                tc = cell._tc; tcPr = tc.get_or_add_tcPr()
                m = OxmlElement("w:tcMar")
                for side in ("top","left","bottom","right"):
                    el = OxmlElement(f"w:{side}")
                    el.set(qn("w:w"), str(pad)); el.set(qn("w:type"),"dxa")
                    m.append(el)
                tcPr.append(m)

            doc = DocxDocument(docx_path)
            for table in doc.tables:
                col_count = max(len(r.cells) for r in table.rows) if table.rows else 0
                tbl = table._tbl; tblPr = tbl.tblPr
                tblW = OxmlElement("w:tblW"); tblW.set(qn("w:w"),"5000"); tblW.set(qn("w:type"),"pct")
                tblPr.append(tblW)
                for ri, row in enumerate(table.rows):
                    for cell in row.cells:
                        _border(cell); _pad(cell)
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.font.size = Pt(10)
                                if ri == 0: run.bold = True
            doc.save(docx_path)
        except Exception as te:
            logger.warning(f"DOCX table formatting skipped: {te}")

        with open(docx_path, "rb") as f:
            docx_bytes = f.read()
        return StreamingResponse(
            iter([docx_bytes]),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="KT_{session_id[:8]}.docx"'},
        )
    except AppException:
        raise
    except Exception as e:
        raise AppException(message=str(e), status_code=500)
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)

import streamlit as st
import uuid
import sys
import os
import json

# Add project root to sys.path
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from datetime import datetime
from app.core.config import settings
from app.core.logger import logger
from app.models.schemas import Session, Topic, Message, TopicKnowledge
from app.services.db_service import db_service
from app.services.ai_engine import ai_engine
from app.services.vector_service import vector_service
from app.services.doc_processor import extract_text_from_file, chunk_text
from app.services.github_service import fetch_repo_content

st.set_page_config(page_title="KT Assistant", layout="wide", initial_sidebar_state="expanded")

# Hide Streamlit's default "Press Enter to submit form" instructions
st.markdown("""
    <style>
    div[data-testid="InputInstructions"] {
        display: none !important;
    }
    </style>
""", unsafe_allow_html=True)

def process_knowledge(text: str):
    """Processes technical knowledge across all topics and updates the session state."""
    all_results = ai_engine.multi_topic_validate_and_score(st.session_state.session, text)
    
    for topic in st.session_state.session.topics:
        if topic.id in all_results:
            data = all_results[topic.id]
            old_score = topic.confidence_score
            topic.knowledge = TopicKnowledge(**data.get("knowledge", {}))
            topic.confidence_score = data.get("confidence_score", 0)
            topic.missing_sections = data.get("missing_sections", [])
            
            if topic.confidence_score != old_score:
                logger.info(f"Topic '{topic.name}' updated. Confidence: {old_score}% -> {topic.confidence_score}%")

            if topic.confidence_score >= settings.KT_CONFIDENCE_THRESHOLD:
                if not topic.is_complete:
                    topic.is_complete = True
                    # Index to Vector DB (RAG Prep)
                    summary_text = json.dumps(topic.knowledge.model_dump(), indent=2)
                    embedding = ai_engine.get_embedding(f"Topic: {topic.name}\nContent: {summary_text}")
                    vector_service.upsert_topic_summary(
                        st.session_state.session_id, 
                        topic.name, 
                        summary_text, 
                        embedding
                    )
    
    st.session_state.session.overall_confidence = int(sum(t.confidence_score for t in st.session_state.session.topics) / len(st.session_state.session.topics))
    db_service.save_session(st.session_state.session)


def index_chunks(chunks: list):
    """Embeds and indexes raw content chunks into Qdrant for Q&A RAG."""
    if not chunks:
        return
    texts = [c["content"] for c in chunks]
    with st.spinner(f"Indexing {len(chunks)} content chunks for Q&A..."):
        embeddings = ai_engine.get_embeddings_batch(texts)
        vector_service.upsert_content_chunks(
            st.session_state.session_id, chunks, embeddings
        )
    logger.info(f"Indexed {len(chunks)} chunks for session {st.session_state.session_id}")

# --- View State & Session Logic ---
if "view" not in st.session_state:
    # If a session ID is provided in the URL, go straight to chat
    url_session_id = st.query_params.get("session_id")
    if url_session_id:
        st.session_state.view = "chat"
    else:
        st.session_state.view = "landing"

# --- Automatic Data Cleanup (6 Hour TTL) ---
if "cleanup_done" not in st.session_state:
    with st.spinner("Performing periodic maintenance..."):
        # 1. Expire sessions older than 6 hours
        expired_ids = db_service.cleanup_expired_sessions(hours=6)
        expired_count = len(expired_ids)
        
        # 2. Get whitelist of all remaining active sessions
        active_ids = db_service.get_all_active_session_ids()
        
        # 3. Purge any vectors NOT in the whitelist (Zombies)
        qdrant_deleted_count = vector_service.purge_zombie_vectors(active_ids)
        
        logger.info(f"Healthcheckup done: {expired_count} expired sessions deleted from Supabase and {qdrant_deleted_count} from qdrant")
        st.session_state.cleanup_done = True

# --- Chat Interface Initialization (Only if in chat view) ---
if st.session_state.view == "chat":
    if "session_id" not in st.session_state:
        url_session_id = st.query_params.get("session_id")
        
        if url_session_id:
            # Try to load existing session from Supabase
            existing_session = db_service.get_session(url_session_id)
            if existing_session:
                st.session_state.session_id = url_session_id
                st.session_state.session = existing_session
                st.session_state.chat_history = db_service.get_messages(url_session_id)
                logger.info(f"Loaded existing session from URL: {url_session_id}")
            else:
                # Redirect or start fresh if ID invalid
                st.session_state.session_id = str(uuid.uuid4())
                st.query_params["session_id"] = st.session_state.session_id
                logger.warning(f"URL session {url_session_id} not found. Starting new.")
        else:
            # No ID anywhere - this shouldn't happen with the current button logic 
            # but we handle it for safety
            st.session_state.session_id = str(uuid.uuid4())
            st.query_params["session_id"] = st.session_state.session_id

    # Initialize Session Object & Chat History if still empty
    if "session" not in st.session_state:
        initial_topics = [
            Topic(id="t1", name="System Overview", missing_sections=["definition", "purpose"]),
            Topic(id="t2", name="Architecture & Data Flow", missing_sections=["inputs / outputs", "monitoring / deployment"]),
            Topic(id="t3", name="Operations & Reliability", missing_sections=["failure cases", "edge cases", "operational steps"])
        ]
        st.session_state.session = Session(id=st.session_state.session_id, topics=initial_topics)
        db_service.save_session(st.session_state.session)

    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []
        greeting = Message(role="assistant", content="Hello! I'm your KT Assistant. Upload a GitHub repository or document from the sidebar, then ask me anything about your codebase and document it.")
        st.session_state.chat_history.append(greeting)
        db_service.save_message(st.session_state.session_id, greeting)

if st.session_state.view == "landing":
    st.markdown("""
        <style>
        .stApp {
            background: #0e1117;
        }
        .main-container {
            max-width: 900px;
            margin: 0 auto;
            padding: 4rem 1rem;
            text-align: center;
        }
        .hero-title {
            font-size: 4rem;
            font-weight: 800;
            background: linear-gradient(90deg, #4facfe 0%, #00f2fe 100%);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            margin-bottom: 1.5rem;
            letter-spacing: -1.5px;
        }
        .hero-desc {
            font-size: 1.4rem;
            color: #8892b0;
            line-height: 1.6;
            margin-bottom: 3rem;
        }
        .highlight-box {
            background: rgba(255, 255, 255, 0.03);
            border: 1px solid rgba(255, 255, 255, 0.1);
            border-radius: 20px;
            padding: 2.5rem;
            margin-bottom: 3rem;
            text-align: left;
        }
        .highlight-item {
            margin-bottom: 1.5rem;
            display: flex;
            align-items: flex-start;
        }
        .highlight-icon {
            font-size: 1.8rem;
            margin-right: 1.2rem;
            margin-top: -4px;
        }
        .highlight-text b {
            color: #4facfe;
            display: block;
            font-size: 1.2rem;
            margin-bottom: 0.2rem;
        }
        .highlight-text p {
            color: #a8b2d1;
            margin: 0;
            font-size: 1.05rem;
        }
        /* Button Styling */
        div.stButton > button {
            background: linear-gradient(90deg, #4facfe 0%, #00f2fe 100%) !important;
            color: white !important;
            font-weight: 700 !important;
            padding: 0.8rem 4rem !important;
            border-radius: 50px !important;
            border: none !important;
            transition: all 0.3s ease !important;
            font-size: 1.3rem !important;
            box-shadow: 0 4px 20px rgba(0, 242, 254, 0.4) !important;
        }
        div.stButton > button:hover {
            transform: scale(1.05) !important;
            box-shadow: 0 10px 30px rgba(0, 242, 254, 0.6) !important;
        }
        </style>
    """, unsafe_allow_html=True)

    with st.container():
        st.markdown('<div class="main-container">', unsafe_allow_html=True)
        st.markdown('<h1 class="hero-title">Knowledge Transfer Assistant</h1>', unsafe_allow_html=True)
        st.markdown('<p class="hero-desc">I am your technical documentation partner. Upload your GitHub repository or documents, and I will instantly analyze your codebase to help you generate comprehensive technical documentation.</p>', unsafe_allow_html=True)
        
        st.markdown("""
            <div class="highlight-box">
                <div class="highlight-item">
                    <span class="highlight-icon">🐙</span>
                    <div class="highlight-text">
                        <b>Automated Codebase Analysis</b>
                        <p>Upload a GitHub repository or documentation files to instantly ingest your entire project context.</p>
                    </div>
                </div>
                <div class="highlight-item">
                    <span class="highlight-icon">💬</span>
                    <div class="highlight-text">
                        <b>Intelligent Q&A</b>
                        <p>Ask questions about your codebase and get accurate answers backed by your project's code and documentation.</p>
                    </div>
                </div>
                <div class="highlight-item">
                    <span class="highlight-icon">📄</span>
                    <div class="highlight-text">
                        <b>Structured Artifacts</b>
                        <p>Automatically track topic coverage and generate a professional technical document ready for your team.</p>
                    </div>
                </div>
            </div>
        """, unsafe_allow_html=True)

        if st.button("🚀 Start KT Session"):
            # Force create a fresh session
            new_id = str(uuid.uuid4())
            st.session_state.session_id = new_id
            st.query_params["session_id"] = new_id
            logger.info(f"User started a fresh KT session: {new_id}")
            
            # Initialize fresh state
            initial_topics = [
                Topic(id="t1", name="System Overview", missing_sections=["definition", "purpose"]),
                Topic(id="t2", name="Architecture & Data Flow", missing_sections=["inputs / outputs", "monitoring / deployment"]),
                Topic(id="t3", name="Operations & Reliability", missing_sections=["failure cases", "edge cases", "operational steps"])
            ]
            st.session_state.session = Session(id=new_id, topics=initial_topics)
            db_service.save_session(st.session_state.session)
            
            st.session_state.chat_history = []
            greeting = Message(role="assistant", content="Hello! I'm your KT Assistant. Upload a GitHub repository or document from the sidebar, then ask me anything about your codebase and document it.")
            st.session_state.chat_history.append(greeting)
            db_service.save_message(new_id, greeting)
            
            st.session_state.view = "chat"
            st.rerun()
            
        st.markdown('</div>', unsafe_allow_html=True)

else:
    # --- Sidebar ---
    with st.sidebar:
        st.title("KT Assistant")
        st.divider()

        tab_progress, tab_data = st.tabs(["📊 KT Status", "📥 Add Context"])

        with tab_progress:
            overall_progress = sum(t.confidence_score for t in st.session_state.session.topics) / len(st.session_state.session.topics)
            st.metric(label="Overall Coverage", value=f"{overall_progress:.1f}%")
            st.caption(f"Recommended **{settings.KT_CONFIDENCE_THRESHOLD}%** coverage across all topics for best results.")
            
            st.markdown("### Topics")
            for i, topic in enumerate(st.session_state.session.topics):
                status_icon = "✅" if topic.confidence_score >= settings.KT_CONFIDENCE_THRESHOLD else ("⏳" if topic.confidence_score > 0 else "📝")
                st.markdown(f"**{status_icon} {topic.name}** ({topic.confidence_score}%)")

            st.divider()



            has_any_content = any(t.confidence_score > 0 for t in st.session_state.session.topics)
            if st.button("Generate Final Document", type="primary", disabled=not has_any_content, use_container_width=True):
                logger.info(f"User triggered final document generation for session: {st.session_state.session_id}")
                with st.spinner("Generating professional KT document..."):
                    summary = ai_engine.generate_final_summary(st.session_state.session)
                    import re
                    clean_summary = re.sub(r'<[^>]+>', '', summary)
                    st.session_state.final_summary = clean_summary
                    if "pdf_bytes" in st.session_state:
                        del st.session_state["pdf_bytes"]
                    st.success("Document generated!")

        with tab_data:
            st.subheader("🐙 GitHub Repository")
            with st.form(key="github_fetch_form", border=False):
                github_url = st.text_input(
                    "GitHub URL",
                    placeholder="https://github.com/owner/repo",
                    label_visibility="collapsed",
                    key="github_url_input",
                    autocomplete="off"
                )
                submitted = st.form_submit_button("🚀 Fetch & Analyse Repo", use_container_width=True)
                
            if submitted:
                if not github_url.strip():
                    st.warning("Please enter a GitHub repository URL first.")
                else:
                    with st.spinner("Fetching repository files from GitHub..."):
                        ingest_result = fetch_repo_content(github_url.strip())
                    if not ingest_result.success:
                        st.error(f"❌ {ingest_result.error}")
                    else:
                        repo_msg = Message(
                            role="user",
                            content=(
                                f"🐙 **GitHub Repository Ingested:** `{ingest_result.owner}/{ingest_result.repo}` "
                                f"(branch: `{ingest_result.branch}`) — "
                                f"{len(ingest_result.files_fetched)} files, "
                                f"{ingest_result.total_chars / 1024:.1f} KB processed."
                            )
                        )
                        st.session_state.chat_history.append(repo_msg)
                        db_service.save_message(st.session_state.session_id, repo_msg)

                        with st.spinner("Analysing repository content across KT topics..."):
                            process_knowledge(ingest_result.aggregated_text)
                        
                        index_chunks(ingest_result.chunks)
                        st.session_state.file_manifest = ingest_result.files_fetched

                        st.success(f"✅ Fetched **{len(ingest_result.files_fetched)} files** from `{ingest_result.owner}/{ingest_result.repo}`.")
                        logger.info(f"GitHub repo ingested: {ingest_result.summary}")
                        st.rerun()

            st.divider()
            st.subheader("📁 File Upload")
            uploaded_file = st.file_uploader("Upload PDF or TXT", type=["pdf", "txt"], label_visibility="collapsed")
            if uploaded_file:
                if "last_uploaded_file" not in st.session_state or st.session_state.last_uploaded_file != uploaded_file.name:
                    with st.spinner(f"Processing {uploaded_file.name}..."):
                        file_bytes = uploaded_file.read()
                        text = extract_text_from_file(file_bytes, uploaded_file.name)
                        if text:
                            doc_msg = Message(role="user", content=f"📄 **Uploaded Document:** {uploaded_file.name}")
                            st.session_state.chat_history.append(doc_msg)
                            db_service.save_message(st.session_state.session_id, doc_msg)
                            
                            process_knowledge(text)
                            
                            file_chunks = chunk_text(text, source_name=uploaded_file.name, chunk_size=settings.CHUNK_SIZE, chunk_overlap=settings.CHUNK_OVERLAP)
                            index_chunks(file_chunks)

                            existing = st.session_state.get("file_manifest", [])
                            if uploaded_file.name not in existing:
                                st.session_state.file_manifest = existing + [uploaded_file.name]
                            
                            st.session_state.last_uploaded_file = uploaded_file.name
                            st.success(f"✅ Processed {uploaded_file.name} for Q&A!")
                            st.rerun()
                        else:
                            st.error("Could not read file content.")

        st.divider()
        if st.button("🗑️ Clear Session Data", type="secondary", use_container_width=True):
            with st.spinner("Clearing data..."):
                db_service.delete_session_data(st.session_state.session_id)
                vector_service.delete_session_vectors(st.session_state.session_id)
                logger.info(f"Session {st.session_state.session_id} cleared by user")
                st.query_params.clear()
                for key in list(st.session_state.keys()):
                    del st.session_state[key]
                st.rerun()

    # --- Main Chat ---
    has_chunks = any(t.confidence_score > 0 for t in st.session_state.session.topics)

    st.header("KT Assistant Chat")

    if not has_chunks:
        st.info("📂 Upload a GitHub repository or document via the sidebar to enable chat.")

    # Display chat messages
    for i, message in enumerate(st.session_state.chat_history):
        with st.chat_message(message.role):
            st.markdown(message.content)

    # --- Q&A Chat ---
    if prompt := st.chat_input("Ask anything about the uploaded project...", disabled=not has_chunks):
        user_msg = Message(role="user", content=prompt)
        st.session_state.chat_history.append(user_msg)
        db_service.save_message(st.session_state.session_id, user_msg)
        with st.chat_message("user"):
            st.markdown(prompt)

        with st.chat_message("assistant"):
            with st.spinner("Analyzing question..."):
                intents, token_stream = ai_engine.route_and_stream(
                    question=prompt,
                    session=st.session_state.session,
                    session_id=st.session_state.session_id,
                    file_manifest=st.session_state.get("file_manifest", []),
                    vector_service=vector_service,
                )

            answer = st.write_stream(token_stream)

            ai_msg = Message(role="assistant", content=answer)
            st.session_state.chat_history.append(ai_msg)
            db_service.save_message(st.session_state.session_id, ai_msg)
            st.rerun()

    # --- Display Summary if generated ---
    if "final_summary" in st.session_state:
        st.divider()
        st.subheader("📄 Final KT Document")

        # --- Coverage chart above the document ---
        import pandas as pd
        topic_names = [t.name for t in st.session_state.session.topics]
        topic_scores = [t.confidence_score for t in st.session_state.session.topics]
        chart_df = pd.DataFrame({"Coverage (%)": topic_scores}, index=topic_names)
        st.caption("**Topic Coverage at Document Generation**")
        st.bar_chart(chart_df, height=180, color="#4facfe")

        st.divider()


        # --- Mermaid-aware document renderer ---
        def render_document_with_mermaid(markdown_text: str):
            """
            Splits the document on mermaid fenced code blocks.
            Renders text sections with st.markdown and diagrams via Mermaid.js.
            """
            import re
            # Split on ```mermaid ... ``` blocks
            pattern = r'(```mermaid\n.*?```)'
            parts = re.split(pattern, markdown_text, flags=re.DOTALL)

            for part in parts:
                if part.startswith("```mermaid"):
                    # Extract the diagram definition
                    diagram_code = part[len("```mermaid\n"):-3].strip()
                    mermaid_html = f"""
                    <div style="background:#1e2530; border-radius:10px; padding:1.5rem; margin:1rem 0;">
                        <div class="mermaid" style="text-align:center;">{diagram_code}</div>
                    </div>
                    <script type="module">
                        import mermaid from 'https://cdn.jsdelivr.net/npm/mermaid@11/dist/mermaid.esm.min.mjs';
                        mermaid.initialize({{
                            startOnLoad: true,
                            theme: 'dark',
                            themeVariables: {{
                                primaryColor: '#4facfe',
                                primaryTextColor: '#e6edf3',
                                lineColor: '#8892b0',
                                background: '#1e2530'
                            }}
                        }});
                    </script>
                    """
                    st.components.v1.html(mermaid_html, height=400, scrolling=True)
                elif part.strip():
                    st.markdown(part)

        render_document_with_mermaid(st.session_state.final_summary)

        # --- PDF & DOCX generation ---
        if "pdf_bytes" not in st.session_state or "docx_bytes" not in st.session_state:
            from markdown_pdf import MarkdownPdf, Section
            import tempfile
            import re
            import subprocess
            import base64
            import json
            import os
            import urllib.request
            import uuid
            import shutil

            # Shared temp dir for locally-downloaded Mermaid diagram images (used by DOCX/Pandoc)
            export_tmp_dir = tempfile.mkdtemp()

            MIN_DIAGRAM_BYTES = 3000  # mermaid.ink error thumbnails are < 1KB; real diagrams > 3KB

            def _download_mermaid_image(code: str) -> tuple[bytes, str] | None:
                """Download a Mermaid diagram from mermaid.ink/img.
                Returns (image_bytes, mime_type) or None on failure."""
                state = {"code": code, "mermaid": {"theme": "default"}}
                b64 = base64.urlsafe_b64encode(json.dumps(state).encode('utf-8')).decode('utf-8')
                img_url = f"https://mermaid.ink/img/{b64}"
                try:
                    req = urllib.request.Request(
                        img_url,
                        headers={'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
                    )
                    with urllib.request.urlopen(req, timeout=15) as resp:
                        data = resp.read()
                    # Detect format by magic bytes
                    if data[:4] == b'\x89PNG':
                        mime = "image/png"
                    elif data[:3] == b'\xff\xd8\xff':
                        # JPEG/JFIF — mermaid.ink returns JPEG by default
                        mime = "image/jpeg"
                    else:
                        logger.warning(f"mermaid.ink returned unexpected format. Code: {code[:80]}")
                        return None
                    # Reject tiny error thumbnails — real diagrams are always > 3KB
                    if len(data) < MIN_DIAGRAM_BYTES:
                        logger.warning(f"mermaid.ink returned tiny image ({len(data)} bytes), likely a render error. Code: {code[:120]}")
                        return None
                    return (data, mime)
                except Exception as dl_err:
                    logger.error(f"Mermaid image download failed: {dl_err}")
                    return None

            def mermaid_to_html_div(match):
                """For PDF: embed diagram as a mermaid div to be rendered by JS in Playwright."""
                code = match.group(1).strip()
                # Wrap in a div that won't be modified by markdown
                return f'<div class="mermaid" style="text-align:center;">\n{code}\n</div>'

            def mermaid_to_local_file(match):
                """For DOCX/Pandoc: save image to local temp file and return a path link."""
                result = _download_mermaid_image(match.group(1).strip())
                if result:
                    img_bytes, mime_type = result
                    ext = "png" if mime_type == "image/png" else "jpg"
                    local_path = os.path.join(export_tmp_dir, f"mermaid_{uuid.uuid4().hex[:8]}.{ext}")
                    with open(local_path, 'wb') as out:
                        out.write(img_bytes)
                    return f"![Diagram]({local_path})"
                return "_[Diagram — could not be rendered]_"

            # PDF uses local browser rendering; DOCX uses local file paths
            pdf_export_markdown = re.sub(
                r'```mermaid\n(.*?)\n```',
                mermaid_to_html_div,
                st.session_state.final_summary,
                flags=re.DOTALL
            )
            docx_export_markdown = re.sub(
                r'```mermaid\n(.*?)\n```',
                mermaid_to_local_file,
                st.session_state.final_summary,
                flags=re.DOTALL
            )

            # 1. Generate PDF with embedded Mermaid images
            try:
                import markdown
                
                # Convert markdown to HTML using standard extensions
                html_body = markdown.markdown(
                    pdf_export_markdown,
                    extensions=['tables', 'fenced_code']
                )
                
                # Wrap in clean CSS mimicking GitHub Markdown for perfect web-standard rendering
                full_html = f"""
                <!DOCTYPE html>
                <html>
                <head>
                <script type="module">
                    import mermaid from 'https://cdn.jsdelivr.net/npm/mermaid@11/dist/mermaid.esm.min.mjs';
                    mermaid.initialize({{
                        startOnLoad: true,
                        theme: 'default'
                    }});
                </script>
                <style>
                body {{ font-family: -apple-system,BlinkMacSystemFont,"Segoe UI",Helvetica,Arial,sans-serif; line-height: 1.6; color: #24292f; margin: 40px auto; max-width: 900px; padding: 20px; }}
                h1, h2, h3, h4 {{ border-bottom: 1px solid #d0d7de; padding-bottom: 0.3em; margin-top: 24px; margin-bottom: 16px; font-weight: 600; page-break-after: avoid; }}
                table {{ border-collapse: collapse; width: 100%; margin: 16px 0; font-size: 14px; page-break-inside: auto; }}
                tr {{ page-break-inside: avoid; page-break-after: auto; }}
                thead {{ display: table-header-group; }}
                th, td {{ border: 1px solid #d0d7de; padding: 8px 13px; text-align: left; }}
                th {{ font-weight: 600; background-color: #f6f8fa; }}
                tr:nth-child(2n) {{ background-color: #f6f8fa; }}
                code {{ background-color: #afb8c133; padding: 0.2em 0.4em; border-radius: 6px; font-family: ui-monospace,SFMono-Regular,SF Mono,Menlo,Consolas,Liberation Mono,monospace; font-size: 85%; }}
                pre {{ background-color: #f6f8fa; padding: 16px; border-radius: 6px; overflow: auto; page-break-inside: avoid; }}
                pre code {{ background-color: transparent; padding: 0; }}
                img {{ max-width: 100%; height: auto; box-sizing: content-box; display: block; margin: 20px auto; page-break-inside: avoid; }}
                blockquote {{ padding: 0 1em; color: #57606a; border-left: .25em solid #d0d7de; margin: 0; }}
                </style>
                </head>
                <body>
                {html_body}
                </body>
                </html>
                """
                
                with tempfile.NamedTemporaryFile(delete=False, suffix=".html") as tmp_html:
                    tmp_html.write(full_html.encode('utf-8'))
                    html_path = tmp_html.name
                    
                pdf_path = html_path.replace(".html", ".pdf")
                
                # Call Playwright wrapper script to avoid asyncio thread conflicts in Streamlit
                script_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "scripts", "generate_pdf.py")
                subprocess.run([sys.executable, script_path, html_path, pdf_path], check=True)
                
                with open(pdf_path, "rb") as f:
                    st.session_state.pdf_bytes = f.read()
                    
                os.unlink(html_path)
                os.unlink(pdf_path)
            except Exception as e:
                logger.error(f"PDF generation failed: {e}")
                st.error(f"Failed to generate PDF: {e}")

            # 2. Generate DOCX (reuse same local images, then Pandoc)
            try:
                with tempfile.NamedTemporaryFile(delete=False, suffix=".md") as tmp_md:
                    tmp_md.write(docx_export_markdown.encode('utf-8'))
                    md_path = tmp_md.name

                docx_path = md_path + ".docx"
                
                # Call pandoc to convert Markdown with remote images into a DOCX
                process = subprocess.run(
                    ["pandoc", md_path, "-o", docx_path],
                    capture_output=True,
                    text=True
                )

                if process.returncode == 0:
                    # Polish tables: full width, even columns, padding, readable font
                    try:
                        from docx import Document as DocxDocument
                        from docx.oxml.ns import qn
                        from docx.oxml import OxmlElement
                        from docx.shared import Pt, Twips

                        def set_cell_border(cell):
                            tc = cell._tc
                            tcPr = tc.get_or_add_tcPr()
                            tcBorders = OxmlElement('w:tcBorders')
                            for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                                tag = OxmlElement(f'w:{edge}')
                                tag.set(qn('w:val'), 'single')
                                tag.set(qn('w:sz'), '6')
                                tag.set(qn('w:space'), '0')
                                tag.set(qn('w:color'), '4A4A4A')
                                tcBorders.append(tag)
                            tcPr.append(tcBorders)

                        def set_cell_padding(cell, pad_twips=120):
                            tc = cell._tc
                            tcPr = tc.get_or_add_tcPr()
                            tcMar = OxmlElement('w:tcMar')
                            for side in ('top', 'left', 'bottom', 'right'):
                                el = OxmlElement(f'w:{side}')
                                el.set(qn('w:w'), str(pad_twips))
                                el.set(qn('w:type'), 'dxa')
                                tcMar.append(el)
                            tcPr.append(tcMar)

                        def set_table_full_width(table, col_count):
                            """Set table to full text-area width and distribute columns evenly."""
                            tbl = table._tbl
                            tblPr = tbl.tblPr
                            # Full width (100%)
                            tblW = OxmlElement('w:tblW')
                            tblW.set(qn('w:w'), '5000')
                            tblW.set(qn('w:type'), 'pct')
                            tblPr.append(tblW)
                            # Enable autofit
                            tblLayout = OxmlElement('w:tblLayout')
                            tblLayout.set(qn('w:type'), 'autofit')
                            tblPr.append(tblLayout)
                            # Even column widths (page text width ≈ 9360 TWIPs for A4)
                            if col_count > 0:
                                col_width = 9360 // col_count
                                for tr in tbl.tr_lst:
                                    for tc in tr.tc_lst:
                                        tcPr = tc.get_or_add_tcPr()
                                        w = OxmlElement('w:tcW')
                                        w.set(qn('w:w'), str(col_width))
                                        w.set(qn('w:type'), 'dxa')
                                        tcPr.append(w)

                        def make_spacer_para(space_before=0, space_after=240):
                            """Create an empty paragraph with spacing for use around tables."""
                            p = OxmlElement('w:p')
                            pPr = OxmlElement('w:pPr')
                            spacing = OxmlElement('w:spacing')
                            spacing.set(qn('w:before'), str(space_before))
                            spacing.set(qn('w:after'), str(space_after))
                            pPr.append(spacing)
                            p.append(pPr)
                            return p

                        doc = DocxDocument(docx_path)

                        # Step 1: Style all tables
                        for i, table in enumerate(doc.tables):
                            col_count = max(len(row.cells) for row in table.rows) if table.rows else 0
                            set_table_full_width(table, col_count)
                            for row_idx, row in enumerate(table.rows):
                                for cell in row.cells:
                                    set_cell_border(cell)
                                    set_cell_padding(cell, pad_twips=120)
                                    for para in cell.paragraphs:
                                        for run in para.runs:
                                            run.font.size = Pt(10)
                                            if row_idx == 0:  # Header row bold
                                                run.bold = True

                        # Step 2: Add spacer paragraphs before and after each table
                        for table in doc.tables:
                            tbl = table._tbl
                            tbl.addprevious(make_spacer_para(space_before=0, space_after=120))
                            tbl.addnext(make_spacer_para(space_before=120, space_after=240))

                        # Step 3: Global paragraph spacing for all body text
                        for para in doc.paragraphs:
                            pPr = para._p.get_or_add_pPr()
                            spacing = pPr.find(qn('w:spacing'))
                            if spacing is None:
                                spacing = OxmlElement('w:spacing')
                                pPr.append(spacing)
                            # Don't override heading spacing — only set on Normal paragraphs
                            if para.style.name.startswith('Heading'):
                                spacing.set(qn('w:before'), '280')
                                spacing.set(qn('w:after'), '80')
                            else:
                                spacing.set(qn('w:after'), '120')

                        doc.save(docx_path)
                        logger.info("DOCX table formatting applied successfully.")
                    except Exception as table_err:
                        logger.warning(f"Failed to format DOCX tables: {table_err}")


                    with open(docx_path, "rb") as f:
                        st.session_state.docx_bytes = f.read()

                else:
                    logger.error(f"Pandoc failed: {process.stderr}")
                    st.error("Failed to generate DOCX file (Pandoc error).")

                # Cleanup
                if os.path.exists(md_path):
                    os.unlink(md_path)
                if os.path.exists(docx_path):
                    os.unlink(docx_path)

            except Exception as e:
                logger.error(f"DOCX generation failed: {e}")
                st.error("Failed to generate DOCX (Pandoc might not be installed).")
            finally:
                # Clean up locally-downloaded Mermaid images
                if os.path.exists(export_tmp_dir):
                    shutil.rmtree(export_tmp_dir)

        # --- Download Buttons ---
        col1, col2 = st.columns(2)
        if "pdf_bytes" in st.session_state:
            with col1:
                st.download_button(
                    label="📥 Download as PDF",
                    data=st.session_state.pdf_bytes,
                    file_name=f"KT_Document_{st.session_state.session_id[:8]}.pdf",
                    mime="application/pdf",
                    use_container_width=True
                )
        if "docx_bytes" in st.session_state:
            with col2:
                st.download_button(
                    label="📥 Download as Word (DOCX)",
                    data=st.session_state.docx_bytes,
                    file_name=f"KT_Document_{st.session_state.session_id[:8]}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
